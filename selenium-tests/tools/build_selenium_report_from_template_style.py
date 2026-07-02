from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Bao_cao_7_2_Selenium_theo_mau_Katalon.docx")


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 13.5), ("Heading 3", 13)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, first_line=True):
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_bullet(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(13)
    r2 = p.add_run(text)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(13)


def add_result_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = True
    headers = ["Nhóm testcase", "Số lượng", "Kết quả chính", "Ghi chú"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
    rows = [
        ["LoginTest", "4", "4 Passed", "Kiểm tra đăng nhập hợp lệ, sai thông tin, bỏ trống dữ liệu và thiếu mật khẩu."],
        ["GuestTourTest, GuestSearchTest", "2", "2 Passed", "Kiểm tra xem chi tiết tour và tìm kiếm tour ở trang public."],
        ["AdminTourTest, AdminUserTest", "4", "4 Passed", "Kiểm tra tìm kiếm/filter tour và tìm kiếm người dùng trong khu vực admin."],
        ["AdminVoucherTest", "1", "1 Failed", "Fail do Selenium không đăng nhập được vào trang admin voucher, vẫn ở /login."],
        ["OperatorTourCreateTest", "1", "1 Failed", "Fail do tài khoản operator không rời khỏi /login khi chuẩn bị vào trang tạo tour."],
        ["GuideProfileTest", "2", "2 Passed", "Kiểm tra thông tin định danh không được sửa và validate số điện thoại quá ngắn."],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11.5)
    doc.add_paragraph()


def main():
    doc = Document()
    style_doc(doc)

    add_heading(doc, "7.2. Phát hiện và báo cáo bug tự động với Selenium", 1)
    add_para(
        doc,
        "Trong phần này, Selenium WebDriver được sử dụng để tự động hóa quá trình kiểm thử giao diện web của hệ thống Chip3Chip. "
        "Khác với kiểm thử thủ công, người kiểm thử không cần trực tiếp ngồi quan sát từng thao tác để kết luận kết quả. "
        "Kịch bản kiểm thử sẽ tự động mở trình duyệt, nhập dữ liệu, thực hiện thao tác trên giao diện và dùng assertion của TestNG "
        "để tự đánh giá trạng thái Passed hoặc Failed. Khi phát sinh lỗi, hệ thống sẽ tự động ghi log, lưu ảnh chụp màn hình và xuất báo cáo phục vụ cho việc phân tích bug."
    )

    add_heading(doc, "7.2.1. Áp dụng vào UC 1", 1)
    add_para(
        doc,
        "UC 1 được áp dụng cho chức năng đăng nhập và một số luồng kiểm thử liên quan đến các vai trò trong hệ thống. "
        "Bộ kiểm thử hiện tại được tổ chức dưới dạng Maven project, sử dụng TestNG suite để chạy tập trung toàn bộ các testcase đã triển khai. "
        "Trong lần thực thi mới nhất, suite all-implemented.xml đã chạy 14 testcase thuộc các nhóm Login, Guest, Admin, Operator và Guide."
    )

    add_heading(doc, "7.2.1.1. Các bước thực hiện", 2)
    add_para(
        doc,
        "Bước 1. Lên kịch bản kiểm thử: Các kịch bản được thiết kế bám sát thao tác thực tế của người dùng trên website. "
        "Đối với chức năng đăng nhập, hệ thống tự động truy cập trang http://localhost:5173/login, nhập email và mật khẩu theo từng bộ dữ liệu, "
        "nhấn nút Đăng nhập, sau đó kiểm tra người dùng có rời khỏi trang /login hay không. Ngoài tài khoản hợp lệ, kịch bản còn kiểm tra các trường hợp "
        "nhập sai thông tin, bỏ trống email và mật khẩu, hoặc chỉ nhập email nhưng không nhập mật khẩu."
    )
    add_para(
        doc,
        "Bước 2. Tổ chức mã nguồn theo Page Object Model: Các thao tác trực tiếp với giao diện không được viết lẫn trong testcase mà được tách vào lớp LoginPage. "
        "Lớp này chịu trách nhiệm định nghĩa locator cho ô email, ô mật khẩu, nút đăng nhập và các thông báo lỗi. Nhờ đó, LoginTest chỉ tập trung mô tả luồng kiểm thử "
        "và điều kiện kiểm tra, còn phần thay đổi giao diện có thể bảo trì tập trung trong Page Object."
    )
    add_caption(doc, "Hình: Cấu trúc LoginTest gọi LoginPage để thao tác với form đăng nhập")
    add_para(
        doc,
        "Bước 3. Chuẩn bị dữ liệu kiểm thử: Dữ liệu đầu vào được lưu trong file CSV tại src/test/resources/test-data/auth-login.csv. "
        "Cách tổ chức này giúp tách biệt dữ liệu khỏi mã nguồn, dễ bổ sung thêm các bộ input như tài khoản hợp lệ, sai thông tin đăng nhập, bỏ trống dữ liệu hoặc thiếu mật khẩu. "
        "Đây là hướng tiếp cận phù hợp khi cần mở rộng kiểm thử theo nhiều bộ dữ liệu."
    )
    add_caption(doc, "Hình: File auth-login.csv chứa các bộ dữ liệu kiểm thử đăng nhập")
    add_para(
        doc,
        "Bước 4. Thiết lập Test Suite và chạy bằng Maven: Các testcase được gom trong file src/test/resources/suites/all-implemented.xml. "
        "File pom.xml đã cấu hình Maven Surefire Plugin để khi chạy lệnh mvn test hoặc double click lifecycle test trong Maven, hệ thống sẽ tự động chạy toàn bộ suite này. "
        "Trong lần chạy lại, suite đã thực thi trên Chrome ở chế độ headless với demo.pause.seconds=0 để giảm thời gian chờ."
    )
    add_caption(doc, "Hình: Cấu hình all-implemented.xml dùng để chạy toàn bộ testcase Selenium")
    add_para(
        doc,
        "Bước 5. Ghi nhận kết quả và báo cáo lỗi: Sau mỗi testcase, ExtentReportListener tiếp nhận kết quả từ TestNG. "
        "Nếu testcase Passed, hệ thống ghi trạng thái thành công vào report. Nếu testcase Failed, hệ thống tự động chụp screenshot, lưu stack trace và ghi thông tin lỗi vào file CSV bug report. "
        "Nhờ đó, kết quả kiểm thử có thể được kiểm tra lại bằng dashboard HTML, console log, file CSV và ảnh lỗi."
    )

    add_heading(doc, "7.2.1.2. Kết quả (log, bug report)", 2)
    add_para(
        doc,
        "Quá trình thực thi: Suite Selenium được chạy lại vào lúc 20:03 ngày 28/06/2026 bằng Maven. "
        "Kết quả chính thức từ TestNG/Surefire ghi nhận tổng cộng 14 testcase, trong đó 12 testcase Passed, 2 testcase Failed, 0 testcase Skipped. "
        "Tổng thời gian chạy suite là khoảng 155.984 giây. Các testcase đăng nhập trong LoginTest đều Passed, xác nhận chức năng đăng nhập với dữ liệu hợp lệ và các trường hợp validation cơ bản hoạt động đúng theo mong đợi."
    )
    add_caption(doc, "Hình: Console Maven hiển thị Tests run: 14, Failures: 2, Errors: 0, Skipped: 0")
    add_result_table(doc)
    add_para(
        doc,
        "Kết quả kiểm tra qua log: File automation.log ghi lại trình tự bắt đầu và kết thúc của từng testcase. "
        "Các dòng log cho thấy LoginTest gồm AUTH_001, AUTH_002, AUTH_003 và AUTH_004 đều đạt trạng thái PASS. "
        "Tuy nhiên, trong quá trình chạy suite tổng, hai testcase AdminVoucherTest.test_ADM_001_search_and_filter_voucher_by_status và "
        "OperatorTourCreateTest.test_OPE_003_required_fields_are_marked_on_create_tour_form bị đánh dấu FAIL do sau bước đăng nhập, trình duyệt vẫn ở URL http://localhost:5173/login."
    )
    add_caption(doc, "Hình: Log Viewer/automation.log ghi nhận trạng thái START, PASS và FAIL của từng testcase")
    add_para(
        doc,
        "Báo cáo bug tự động: File target/reports/bug_report.csv đã ghi lại thông tin testcase, tên method, trạng thái, nội dung lỗi và đường dẫn screenshot. "
        "Đối với hai testcase fail, lỗi được ghi nhận là TimeoutException với điều kiện chờ URL không còn chứa /login nhưng kết quả thực tế vẫn là trang login. "
        "Điều này cho thấy hệ thống automation không chỉ báo testcase fail mà còn mô tả rõ vị trí lỗi và trạng thái thực tế tại thời điểm thất bại."
    )
    add_para(
        doc,
        "Bằng chứng hình ảnh: Khi testcase thất bại, screenshot được tự động lưu vào thư mục target/screenshots. "
        "Hai ảnh lỗi mới nhất được tạo trong lần chạy này là AdminVoucherTest_test_ADM_001_search_and_filter_voucher_by_status_20260628_200521.png và "
        "OperatorTourCreateTest_test_OPE_003_required_fields_are_marked_on_create_tour_form_20260628_200551.png. "
        "Các ảnh này giúp người kiểm thử hoặc giảng viên xem lại giao diện thực tế tại thời điểm lỗi mà không cần chạy lại testcase ngay lập tức."
    )
    add_caption(doc, "Hình: Thư mục target/screenshots chứa ảnh màn hình khi testcase fail")
    add_para(
        doc,
        "Báo cáo tổng hợp: Ngoài console và CSV, hệ thống còn sinh ExtentReports tại target/extent-report/ExtentReport.html và báo cáo TestNG tại target/surefire-reports/index.html. "
        "ExtentReports phù hợp để trình bày demo vì có giao diện dashboard, thể hiện danh sách testcase, trạng thái pass/fail và log theo từng bước. "
        "TestNG report và file XML trong surefire-reports được dùng để đối chiếu số liệu chính thức của lần chạy."
    )
    add_caption(doc, "Hình: ExtentReports Dashboard hiển thị danh sách testcase và trạng thái thực thi")

    add_heading(doc, "7.2.1.3. Nhận xét", 2)
    add_para(doc, "Ưu điểm:", first_line=False)
    add_bullet(doc, "Tự động phát hiện lỗi", "Selenium kết hợp TestNG assertion có thể tự kết luận testcase pass hoặc fail mà không cần người kiểm thử quan sát thủ công từng màn hình.")
    add_bullet(doc, "Minh chứng rõ ràng", "Khi testcase fail, hệ thống tự động lưu screenshot, stack trace, log và báo cáo HTML/CSV, hỗ trợ quá trình báo cáo bug.")
    add_bullet(doc, "Dễ mở rộng dữ liệu", "Các bộ input đăng nhập được tách ra file CSV, giúp bổ sung thêm trường hợp kiểm thử mà không cần viết lại toàn bộ kịch bản.")
    add_bullet(doc, "Dễ bảo trì", "Page Object Model giúp gom locator và thao tác giao diện vào LoginPage, giảm trùng lặp code trong testcase.")
    add_para(doc, "Nhược điểm:", first_line=False)
    add_bullet(doc, "Phụ thuộc môi trường chạy", "Nếu backend chưa ổn định, tài khoản bị giới hạn đăng nhập hoặc phiên đăng nhập hết hạn, testcase có thể fail dù logic code automation không sai.")
    add_bullet(doc, "Cần bảo trì locator", "Khi giao diện website thay đổi id, class hoặc cấu trúc DOM, locator trong Page Object cần được cập nhật kịp thời.")
    add_bullet(doc, "Có thể phát sinh lỗi không ổn định", "Hai testcase fail trong lần chạy cho thấy quá trình đăng nhập ở một số role có thể bị timeout hoặc không rời khỏi trang login, cần kiểm tra thêm nguyên nhân ở backend/session hoặc tài khoản test.")
    add_para(
        doc,
        "Tổng kết: Việc áp dụng Selenium cho UC 1 đã chứng minh được workflow phát hiện và báo cáo bug tự động. "
        "Hệ thống có thể chạy suite, thao tác giao diện, tự đánh giá kết quả, sinh report và lưu bằng chứng lỗi. "
        "Trong lần chạy mới nhất, các testcase đăng nhập cốt lõi đều Passed, đồng thời automation cũng phát hiện được hai lỗi liên quan đến luồng đăng nhập trước khi truy cập chức năng Admin Voucher và Operator Create Tour. "
        "Đây là cơ sở để tiếp tục mở rộng kiểm thử tự động cho các use case khác của hệ thống Chip3Chip."
    )

    doc.save(OUT)


if __name__ == "__main__":
    main()
