from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Bao_cao_7_2_Phat_hien_va_bao_cao_bug_tu_dong_Selenium.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181) if level < 3 else RGBColor(31, 77, 120)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_info_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="1F4D78")
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()
    return table


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def main():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("7.2. Phát hiện và báo cáo bug tự động với Selenium")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Áp dụng cho UC 1: Kiểm thử chức năng đăng nhập")
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "7.2. Phát hiện và báo cáo bug tự động với Selenium", 1)
    doc.add_paragraph(
        "Trong hệ thống kiểm thử hiện tại, Selenium WebDriver được dùng để thao tác trực tiếp trên giao diện website "
        "như người dùng thật. TestNG chịu trách nhiệm tổ chức testcase và đưa ra kết luận Pass/Fail thông qua assertion. "
        "Khi testcase thất bại, listener sẽ tự động ghi nhận thông tin lỗi, lưu ảnh chụp màn hình và tổng hợp kết quả "
        "vào báo cáo."
    )

    add_bullets(doc, [
        "Website kiểm thử: http://localhost:5173/",
        "Công cụ chính: Java, Selenium WebDriver, TestNG, ExtentReports.",
        "Bằng chứng sau khi chạy: log từng bước, ExtentReport HTML, file CSV bug report và screenshot khi fail.",
    ])

    add_info_table(
        doc,
        ["Thành phần", "Vai trò trong workflow"],
        [
            ["Test data CSV", "Lưu các bộ dữ liệu đăng nhập như tài khoản hợp lệ, sai thông tin, bỏ trống dữ liệu."],
            ["LoginTest", "Chứa các testcase của UC 1 và gọi Page Object để thao tác với website."],
            ["LoginPage", "Đóng gói locator và thao tác trên form đăng nhập: nhập email, password, click Login, kiểm tra lỗi."],
            ["ExtentReportListener", "Lắng nghe kết quả TestNG, ghi log, chụp screenshot khi fail và sinh report."],
        ],
        [1.65, 4.75],
    )

    add_heading(doc, "7.2.1. Áp dụng vào UC 1", 1)
    doc.add_paragraph(
        "UC 1 được áp dụng cho chức năng đăng nhập. Mục tiêu là kiểm tra hệ thống có cho phép người dùng đăng nhập "
        "với tài khoản hợp lệ hay không, đồng thời phát hiện các lỗi khi người dùng nhập sai hoặc thiếu dữ liệu."
    )

    add_heading(doc, "7.2.1.1. Các bước thực hiện", 2)
    add_numbered(doc, [
        "Chuẩn bị dữ liệu test trong file src/test/resources/test-data/auth-login.csv.",
        "Maven chạy TestNG suite được cấu hình trong pom.xml.",
        "TestNG gọi LoginTest và thực thi lần lượt các testcase của UC 1.",
        "Selenium mở trình duyệt Chrome và truy cập trang http://localhost:5173/login.",
        "LoginPage tìm phần tử bằng locator, nhập username/password và nhấn nút Đăng nhập.",
        "TestNG dùng assertion để kiểm tra kết quả thực tế so với kết quả mong đợi.",
        "ExtentReportListener ghi log, cập nhật trạng thái Pass/Fail và lưu screenshot nếu testcase fail.",
    ])

    add_info_table(
        doc,
        ["Mã testcase", "Dữ liệu/điều kiện", "Kết quả mong đợi"],
        [
            ["AUTH_001", "Email và mật khẩu admin hợp lệ", "Đăng nhập thành công, URL không còn ở /login."],
            ["AUTH_002", "Thông tin đăng nhập không hợp lệ", "Vẫn ở trang login và hiển thị thông báo lỗi."],
            ["AUTH_003", "Bỏ trống email và password", "Hiển thị validate bắt buộc nhập email và password."],
            ["AUTH_004", "Nhập email nhưng bỏ trống password", "Hiển thị validate bắt buộc nhập password."],
        ],
        [1.25, 2.55, 2.6],
    )

    add_heading(doc, "7.2.1.2. Kết quả (log, bug report)", 2)
    doc.add_paragraph(
        "Sau khi chạy testcase, hệ thống tự động tạo các loại bằng chứng kiểm thử. Nhờ đó người kiểm thử không cần "
        "ngồi quan sát thủ công từng bước trên màn hình mà có thể xem kết quả tổng hợp sau khi suite kết thúc."
    )

    add_info_table(
        doc,
        ["Loại kết quả", "Vị trí", "Ý nghĩa"],
        [
            ["Console log", "Màn hình Run/Maven trong IDE", "Hiển thị số testcase đã chạy, số pass/fail và lỗi tổng quát."],
            ["ExtentReports", "target/extent-report/ExtentReport.html", "Dashboard HTML có danh sách testcase, trạng thái và log từng bước."],
            ["Bug report CSV", "target/reports/bug_report.csv", "Bảng tổng hợp testcase, method, status, lỗi và đường dẫn screenshot."],
            ["Screenshot lỗi", "target/screenshots", "Ảnh chụp giao diện tại thời điểm testcase thất bại, dùng làm bằng chứng bug."],
            ["TestNG report", "target/surefire-reports", "Báo cáo mặc định của Maven/TestNG để đối chiếu kết quả chạy."],
        ],
        [1.35, 2.55, 2.5],
    )

    doc.add_paragraph(
        "Ví dụ: nếu testcase đăng nhập thành công nhưng hệ thống vẫn giữ người dùng ở trang /login, assertion sẽ fail. "
        "Listener sẽ ghi trạng thái FAIL vào report, lưu stack trace và chụp màn hình tại thời điểm lỗi để phục vụ báo cáo bug."
    )

    add_heading(doc, "7.2.1.3. Nhận xét", 2)
    add_bullets(doc, [
        "Automation test giúp phát hiện lỗi nhanh hơn so với kiểm thử thủ công, đặc biệt khi phải chạy lại nhiều bộ dữ liệu.",
        "Report tự động giúp xác định testcase nào pass, testcase nào fail, fail ở bước nào và có bằng chứng đi kèm.",
        "Việc tách LoginTest và LoginPage giúp code dễ bảo trì: nếu giao diện login thay đổi, chỉ cần cập nhật locator trong Page Object.",
        "Đối với tài khoản thật, cần hạn chế chạy nhiều testcase sai mật khẩu liên tục để tránh cơ chế khóa tạm thời hoặc rate limit.",
        "Khi trình bày demo, ExtentReports phù hợp để minh họa dashboard kết quả, còn screenshot fail và CSV phù hợp để chứng minh khả năng báo cáo bug tự động.",
    ])

    conclusion = doc.add_paragraph()
    conclusion.paragraph_format.space_before = Pt(8)
    conclusion.paragraph_format.space_after = Pt(0)
    run = conclusion.add_run(
        "Kết luận: Với UC 1, hệ thống Selenium đã có workflow tự động từ dữ liệu test, thao tác giao diện, kiểm tra assertion "
        "đến sinh report. Đây là cơ sở để mở rộng sang các use case khác trong website."
    )
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 77, 120)

    doc.save(OUT)


if __name__ == "__main__":
    main()
